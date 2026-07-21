import streamlit as st
import google.generativeai as genai
from fpdf import FPDF
import os
import requests
import time
import hashlib
import datetime
from PIL import Image
import io
import json
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# Опциональные библиотеки
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import qrcode
    HAS_QR = True
except ImportError:
    HAS_QR = False

# ==========================================
# 1. LUXURY ENTERPRISE STYLING (CSS)
# ==========================================
st.set_page_config(
    page_title="Global Asset Auditor | Ultimate Platform", 
    page_icon="🏢", 
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .stDeployButton {display:none;}
    
    .stApp {
        background-color: #0B0C0E;
        color: #E2E8F0;
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    }
    
    h1, h2, h3 {
        color: #D4AF37 !important;
        font-weight: 700 !important;
        letter-spacing: -0.5px;
    }
    
    div[data-testid="stExpander"] {
        border: 1px solid #23252D !important;
        background-color: #141519 !important;
        border-radius: 10px !important;
    }
    
    .stButton>button {
        background: linear-gradient(135deg, #E6C200 0%, #B8860B 100%);
        color: #0A0A0C !important;
        font-weight: 800 !important;
        font-size: 16px !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 14px 28px !important;
        width: 100%;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
        box-shadow: 0 4px 20px rgba(212, 175, 55, 0.25) !important;
    }
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(212, 175, 55, 0.45) !important;
        background: linear-gradient(135deg, #F0D000 0%, #C5930C 100%);
    }
    
    .stTextInput>div>div>input, .stTextArea>div>div>textarea, .stSelectbox>div>div {
        background-color: #141519 !important;
        color: #FFFFFF !important;
        border: 1px solid #2A2C36 !important;
        border-radius: 8px !important;
    }
    .stTextInput>div>div>input:focus, .stTextArea>div>div>textarea:focus {
        border-color: #D4AF37 !important;
        box-shadow: 0 0 10px rgba(212, 175, 55, 0.2) !important;
    }
    
    section[data-testid="stSidebar"] {
        background-color: #0E0F12 !important;
        border-right: 1px solid #1E2028;
    }
    </style>
""", unsafe_allow_html=True)

# ==========================================
# 2. РЕЕСТР И ИСТОРИЯ (DB MANAGER)
# ==========================================
HISTORY_FILE = "audit_history.json"

def load_history():
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []
    return []

def save_to_history(record):
    history = load_history()
    history.insert(0, record)
    with open(HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=2)

# ==========================================
# 3. ШРИФТЫ И QR-КОДЫ
# ==========================================
FONT_PATH = "DejaVuSans.ttf"

def get_validated_font():
    if os.path.exists(FONT_PATH) and os.path.getsize(FONT_PATH) > 100000:
        try:
            with open(FONT_PATH, "rb") as f:
                if f.read(4) in [b'\x00\x01\x00\x00', b'OTTO', b'true']:
                    return FONT_PATH
        except Exception:
            pass

    urls = [
        "https://cdn.jsdelivr.net/npm/dejavu-fonts-ttf@2.37.0/ttf/DejaVuSans.ttf",
        "https://raw.githubusercontent.com/fpdf2/fpdf2/master/test/fonts/DejaVuSans.ttf",
        "https://raw.githubusercontent.com/matomo-org/travis-scripts/master/fonts/DejaVuSans.ttf"
    ]
    
    for url in urls:
        try:
            res = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
            if res.status_code == 200 and len(res.content) > 100000:
                if res.content[:4] in [b'\x00\x01\x00\x00', b'OTTO', b'true']:
                    with open(FONT_PATH, "wb") as f:
                        f.write(res.content)
                    return FONT_PATH
        except Exception:
            continue
    return None

def generate_qr_code(data_str):
    if not HAS_QR:
        return None
    qr = qrcode.QRCode(version=1, box_size=4, border=1)
    qr.add_data(data_str)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    path = "temp_qr.png"
    img.save(path)
    return path

# ==========================================
# 4. ДВИЖОК ЭКСПОРТА (PDF & WORD)
# ==========================================
class UltimatePDF(FPDF):
    def __init__(self, company_name="GLOBAL ASSET AUDITOR", doc_hash="", logo_path=None, use_dejavu=True):
        super().__init__()
        self.company_name = company_name
        self.doc_hash = doc_hash
        self.logo_path = logo_path
        self.use_dejavu = use_dejavu

    def header(self):
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                self.image(self.logo_path, x=10, y=8, w=22)
            except Exception:
                pass
            
        if self.use_dejavu:
            self.set_font("DejaVu", "", 8)
        else:
            self.set_font("Helvetica", "B", 8)
            
        self.set_text_color(140, 140, 140)
        self.cell(0, 5, f"{self.company_name.upper()} | VERIFIED AUDIT REPORT", 0, 1, "R")
        self.set_draw_color(212, 175, 55)
        self.set_line_width(0.6)
        self.line(10, 16, 200, 16)
        self.ln(7)

    def footer(self):
        self.set_y(-16)
        if self.use_dejavu:
            self.set_font("DejaVu", "", 7)
        else:
            self.set_font("Helvetica", "", 7)
        self.set_text_color(120, 120, 120)
        self.cell(130, 4, f"Page {self.page_no()} | SHA-256 Hash: {self.doc_hash[:20]}...", 0, 0, "L")
        self.cell(60, 4, "Official Audit Document", 0, 1, "R")

def build_docx_report(company_name, inspector_name, object_name, report_text, logo_path=None, image_path=None):
    if not HAS_DOCX:
        return None
    
    doc = docx.Document()
    
    if logo_path and os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.5))
        
    p = doc.add_paragraph()
    run = p.add_run(f"{company_name.upper()} — AUDIT REPORT")
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(184, 134, 11)
    
    doc.add_paragraph(f"Lead Auditor: {inspector_name} | Object: {object_name}")
    doc.add_paragraph("="*50)
    
    for line in report_text.split('\n'):
        if line.startswith('#'):
            clean_line = line.replace('#', '').strip()
            p = doc.add_paragraph()
            r = p.add_run(clean_line)
            r.font.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = RGBColor(184, 134, 11)
        else:
            doc.add_paragraph(line)
            
    if image_path and os.path.exists(image_path):
        doc.add_page_break()
        doc.add_heading('APPENDIX 1: DEFECT PHOTO FIXATION', level=1)
        doc.add_picture(image_path, width=Inches(5.5))
        
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# 5. МОДУЛЬ EMAIL-ОТПРАВКИ
# ==========================================
def send_audit_email(smtp_server, smtp_port, sender_email, sender_password, recipient_email, subject, body, attachment_bytes, filename):
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain', 'utf-8'))
    
    part = MIMEBase('application', 'octet-stream')
    part.set_payload(attachment_bytes)
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
    msg.attach(part)
    
    server = smtplib.SMTP(smtp_server, int(smtp_port))
    server.starttls()
    server.login(sender_email, sender_password)
    server.send_message(msg)
    server.quit()

# ==========================================
# 6. ИИ-ЯДРО (GEMINI AUDIT ENGINE)
# ==========================================
def run_ai_audit(api_key, prompt, image=None):
    genai.configure(api_key=api_key)
    
    try:
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
    except Exception as e:
        raise Exception(f"Ошибка соединения с API: {e}")
        
    preferred_models = ['models/gemini-1.5-flash', 'models/gemini-1.5-pro', 'models/gemini-2.0-flash']
    candidates = [m for m in preferred_models if m in available_models] + [m for m in available_models if m not in preferred_models]
    
    last_err = None
    for model_name in candidates:
        try:
            model = genai.GenerativeModel(model_name)
            content = [prompt, image] if image else [prompt]
            response = model.generate_content(content)
            return response.text, model_name.replace("models/", "")
        except Exception as e:
            last_err = e
            time.sleep(1)
            continue
            
    raise last_err if last_err else Exception("Не удалось обратиться к ИИ-моделям.")

# ==========================================
# 7. ИНТЕРФЕЙС И ЛОГИКА
# ==========================================
st.sidebar.markdown("### 🏢 Auditor Control Panel")
api_key = st.sidebar.text_input("Google Gemini API Key", type="password", help="Введите ваш ключ Google AI Studio")

st.sidebar.markdown("---")
st.sidebar.markdown("### 🎨 Брендинг и Настройки")

# Загрузка логотипа
uploaded_logo = st.sidebar.file_uploader("Логотип компании (PNG/JPG):", type=["png", "jpg", "jpeg"])
temp_logo_path = None
if uploaded_logo:
    logo_img = Image.open(uploaded_logo)
    temp_logo_path = "temp_company_logo.png"
    logo_img.save(temp_logo_path)
    st.sidebar.image(logo_img, caption="Загруженный логотип", width=120)

company_name = st.sidebar.text_input("Название компании:", value="Global Asset Solutions")
inspector_name = st.sidebar.text_input("ФИО Старшего аудитора:", value="Сабит Фатизаде")
doc_lang = st.sidebar.selectbox("Язык акта / Language:", ["Русский", "Azərbaycan dili", "English"])

# Главный заголовок
st.title("🏢 Global Asset Auditor")
st.markdown("##### *Платформа автоматизированного технико-юридического аудита и оценки недвижимости*")
st.markdown("---")

tab_new, tab_history, tab_email = st.tabs(["🚀 Новый аудит", "📜 Реестр и Архив", "📧 Отправка по Email"])

# ------------------------------------------
# ВКЛАДКА 1: НОВЫЙ АУДИТ
# ------------------------------------------
with tab_new:
    if api_key:
        st.sidebar.success("🔑 API Ключ подключен")
        
        col_main1, col_main2 = st.columns([1, 1])
        
        with col_main1:
            st.markdown("#### 1. Основные сведения")
            object_name = st.text_input("Объект / Адрес:", value="Офис №1402, БЦ 'Port Baku Tower', г. Баку")
            audit_type = st.selectbox("Цель аудита:", [
                "Прием-передача объекта недвижимости", 
                "Плановая инвентаризация и дефектовка", 
                "Предпродажный экспресс-аудит для инвестора"
            ])
            
        with col_main2:
            st.markdown("#### 2. Визуальный контроль (Vision AI)")
            uploaded_image = st.file_uploader("Загрузить фото дефекта:", type=["jpg", "png", "jpeg"])
            image_obj = None
            temp_img_path = None
            if uploaded_image:
                image_obj = Image.open(uploaded_image)
                temp_img_path = "temp_audit_defect.jpg"
                image_obj.convert('RGB').save(temp_img_path)
                st.image(image_obj, caption="Снимок для Vision AI и вшивания в PDF", use_container_width=True)

        st.markdown("#### 3. Техническое описание и выявленные замечания")
        user_input = st.text_area("Вводные данные аудитора:", height=130, 
                                  value="Отделка: Микротрещины на гипсокартоне (ресепшн), следы протечек на Armstrong (сектор B), царапины на ламинате. HVAC: Шум кондиционера, требуется чистка фильтров. Электрика: Щит №2 без маркировки. Сантехника: Шатается смеситель. Ресурсы: Э/Э 48512 кВт·ч, ХВ 412 м³, ГВ 189 м³. Ключи: 4 главных, 10 RFID карт, 2 пульта климата.")

        if st.button("🚀 Сформировать Комплексный Акт"):
            if object_name and user_input:
                with st.spinner('ИИ-Аудитор формирует юридический Акт и верстает документы...'):
                    
                    lang_instruction = {
                        "Русский": "Составь документ ПОЛНОСТЬЮ НА РУССКОМ ЯЗЫКЕ.",
                        "Azərbaycan dili": "Sənədi TAMAMİLƏ AZƏRBAYCAN DİLİNDƏ tərtib et. Bütün başlıqlar, cədvəllər və terminlər Azərbaycan dilində olmalıdır.",
                        "English": "Draft the document ENTIRELY IN ENGLISH. All headings, tables, legal terms must be in English."
                    }[doc_lang]
                    
                    full_prompt = f"""
Ты — Senior Asset Auditor и международный юрист.
Составь 'ОФИЦИАЛЬНЫЙ АКТ ТЕХНИЧЕСКОГО АУДИТА И ПРИЕМА-ПЕРЕДАЧИ ОБЪЕКТА'.
ЯЗЫКОВОЕ ТРЕБОВАНИЕ: {lang_instruction}

ИСПОЛНИТЕЛИ И ОБЪЕКТ:
- Аудиторская компания: '{company_name}'
- Старший аудитор: '{inspector_name}'
- Объект аудита: '{object_name}'
- Категория: {audit_type}
- Детальные вводные: {user_input}
{"- Проанализируй прикрепленное фото, определи дефекты и внеси их в соответствующую секцию." if uploaded_image else ""}

ОБЯЗАТЕЛЬНАЯ СТРУКТУРА ДОКУМЕНТА:
1. ЮРИДИЧЕСКАЯ ПРЕАМБУЛА (Место, дата, стороны, основание проверки).
2. ТЕХНИЧЕСКИЙ ПАСПОРТ И ПРИБОРЫ УЧЕТА (Ключи, карты доступа, показания электро/водосчетчиков).
3. ДЕФЕКТОВОЧНАЯ ВЕДОМОСТЬ (Покатегорийный разбор: Конструкции, HVAC/Инженерия, Сантехника, Отделка. Для каждого дефекта укажи КРИТИЧНОСТЬ: Высокая/Средняя/Низкая и рекомендации).
4. ЮРИДИЧЕСКИЙ АНАЛИЗ И ОЦЕНКА РИСКОВ (Фиксация условия 'As Is' / 'Как есть', ограничение ответственности).
5. РЕГЛАМЕНТ И СРОКИ УСТРАНЕНИЯ НЕДОСТАТКОВ (Порядок компенсаций и устранения).
6. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН (Графы для подписей Сторон и Инспектора: {inspector_name}).

Стиль: строго официальный, юридически бескомпромиссный. Выдай готовый документ в формате Markdown без вступительных реплик.
"""
                    try:
                        report_text, used_model = run_ai_audit(api_key, full_prompt, image_obj)
                        
                        st.success(f"✅ Документ успешно сформирован (Язык: {doc_lang} | Модель: {used_model})")
                        
                        # Предпросмотр
                        with st.expander("📄 Просмотр сформированного документа", expanded=True):
                            st.markdown(report_text)
                        
                        # Расчет хэша
                        doc_hash = hashlib.sha256(report_text.encode('utf-8')).hexdigest()
                        qr_path = generate_qr_code(f"VERIFIED AUDIT ACT\nHash: {doc_hash}\nAuditor: {inspector_name}")
                        
                        # Сборка PDF
                        font_file = get_validated_font()
                        use_dejavu = font_file is not None
                        
                        pdf = UltimatePDF(company_name=company_name, doc_hash=doc_hash, logo_path=temp_logo_path, use_dejavu=use_dejavu)
                        
                        if use_dejavu:
                            pdf.add_font("DejaVu", "", font_file)
                        
                        pdf.add_page()
                        
                        if use_dejavu:
                            pdf.set_font("DejaVu", size=10)
                        else:
                            pdf.set_font("Helvetica", size=10)
                            
                        pdf.set_text_color(30, 30, 30)
                        clean_text = report_text.replace('**', '').replace('*', '-').replace('#', '')
                        pdf.multi_cell(0, 6, text=clean_text)
                        
                        # Приложение с фото
                        if temp_img_path and os.path.exists(temp_img_path):
                            pdf.add_page()
                            if use_dejavu:
                                pdf.set_font("DejaVu", size=12)
                            else:
                                pdf.set_font("Helvetica", "B", 12)
                            pdf.set_text_color(184, 134, 11)
                            pdf.cell(0, 10, "APPENDIX 1: DEFECT PHOTO FIXATION", 0, 1, "L")
                            pdf.ln(5)
                            pdf.image(temp_img_path, x=15, y=30, w=180)
                            
                        # QR-код
                        if qr_path and os.path.exists(qr_path):
                            pdf.image(qr_path, x=165, y=240, w=30)
                        
                        pdf_bytes = bytes(pdf.output())
                        
                        # Генерация Word
                        docx_bytes = build_docx_report(company_name, inspector_name, object_name, report_text, temp_logo_path, temp_img_path)
                        
                        # Сохранение в сессию для Email
                        st.session_state['last_pdf'] = pdf_bytes
                        st.session_state['last_filename'] = f"Audit_Act_{object_name[:10].replace(' ', '_')}.pdf"
                        
                        # Сохранение в историю (База данных)
                        save_to_history({
                            "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                            "object_name": object_name,
                            "company": company_name,
                            "auditor": inspector_name,
                            "lang": doc_lang,
                            "hash": doc_hash,
                            "text": report_text
                        })
                        
                        # Кнопки скачивания
                        st.markdown("### 📥 Скачать пакет документов:")
                        col_dl1, col_dl2 = st.columns(2)
                        
                        with col_dl1:
                            st.download_button(
                                label="⬇️ Скачать PDF (с QR, Логотипом и Фото)",
                                data=pdf_bytes,
                                file_name=f"Audit_Act_{object_name[:10].replace(' ', '_')}.pdf",
                                mime="application/pdf"
                            )
                            
                        with col_dl2:
                            if docx_bytes:
                                st.download_button(
                                    label="📝 Скачать Редактируемый Word (.docx)",
                                    data=docx_bytes,
                                    file_name=f"Audit_Act_{object_name[:10].replace(' ', '_')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                                
                    except Exception as err:
                        if "429" in str(err) or "Quota" in str(err):
                            st.error("⏳ **Превышен лимит запросов API.** Подождите 30 секунд или используйте другой ключ.")
                        else:
                            st.error(f"❌ Ошибка системы: {err}")
            else:
                st.warning("⚠️ Пожалуйста, укажите объект и описание состояния.")
    else:
        st.info("👈 Введите ваш Google Gemini API Key в меню слева для активации платформы.")

# ------------------------------------------
# ВКЛАДКА 2: РЕЕСТР И АРХИВ
# ------------------------------------------
with tab_history:
    st.markdown("#### 📜 Электронный реестр проведенных аудитов")
    records = load_history()
    
    if records:
        for idx, rec in enumerate(records):
            with st.expander(f"📌 [{rec['timestamp']}] {rec['object_name']} ({rec['lang']})"):
                st.write(f"**Аудиторская компания:** {rec['company']} | **Инспектор:** {rec['auditor']}")
                st.write(f"**SHA-256 Хэш:** `{rec['hash']}`")
                st.markdown("---")
                st.markdown(rec['text'])
    else:
        st.info("Архив аудитов пока пуст. Сформируйте ваш первый документ!")

# ------------------------------------------
# ВКЛАДКА 3: ОТПРАВКА ПО EMAIL
# ------------------------------------------
with tab_email:
    st.markdown("#### 📧 Рассылка отчетов контрагентам")
    
    if 'last_pdf' in st.session_state:
        col_em1, col_em2 = st.columns(2)
        
        with col_em1:
            smtp_server = st.text_input("SMTP Сервер:", value="smtp.gmail.com")
            smtp_port = st.text_input("SMTP Порт:", value="587")
            sender_email = st.text_input("Email отправителя:", placeholder="your_email@gmail.com")
            sender_password = st.text_input("Пароль приложения SMTP:", type="password")
            
        with col_em2:
            recipient_email = st.text_input("Email получателя (Клиент / Юрист):")
            email_subject = st.text_input("Тема письма:", value=f"Официальный Акт Технического Аудита — Global Asset Solutions")
            email_body = st.text_area("Текст письма:", value="Здравствуйте!\n\nНаправляем вам официально сформированный и подлинный Акт технического аудита и приема-передачи объекта.\n\nДокумент защищен SHA-256 хэшем и содержит QR-код верификации в приложенном PDF-файле.")
            
        if st.button("📤 Отправить Акт по Email"):
            if recipient_email and sender_email and sender_password:
                try:
                    with st.spinner("Отправка письма с вложением..."):
                        send_audit_email(
                            smtp_server, smtp_port, sender_email, sender_password, 
                            recipient_email, email_subject, email_body, 
                            st.session_state['last_pdf'], st.session_state['last_filename']
                        )
                    st.success(f"✅ Документ успешно отправлен на адрес: {recipient_email}")
                except Exception as e:
                    st.error(f"❌ Ошибка отправки: {e}")
            else:
                st.warning("⚠️ Заполните все поля почтового сервера и адреса получателя.")
    else:
        st.info("💡 Сначала сформируйте документ во вкладке '🚀 Новый аудит'.")
